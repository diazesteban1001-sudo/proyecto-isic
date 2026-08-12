#!/usr/bin/env python3
"""
md_a_docx.py — convierte el borrador en Markdown al entregable .docx.

Fase 3 de la skill sintesis-consultoria. Se corre SOLO después de que
verificar_trazabilidad.py esté limpio: convertir antes es invertir
tiempo en formato sobre un texto que todavía puede cambiar.

Uso:
    python md_a_docx.py --entrada informe/borrador.md \
                        --salida informe/informe-final.docx
"""

import argparse
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# **negrita**, *cursiva*, `código` — se parte la línea en trozos y cada
# trozo conserva su marca para aplicarla como formato de run, no como
# texto literal con asteriscos.
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\*[^*]+?\*)")


def escribir_runs(parrafo, texto, negrita=False, cursiva=False, codigo=False):
    """Escribe el texto aplicando el formato inline como formato de run.

    Recursiva a propósito: el informe anida marcas —**`columna` en negrita**—
    y una pasada plana deja los acentos graves como texto literal dentro del
    trozo en negrita."""
    trozos = INLINE.split(texto)
    for trozo in trozos:
        if not trozo:
            continue
        if len(trozos) > 1 and INLINE.fullmatch(trozo):
            if trozo.startswith("**"):
                escribir_runs(parrafo, trozo[2:-2], True, cursiva, codigo)
            elif trozo.startswith("`"):
                escribir_runs(parrafo, trozo[1:-1], negrita, cursiva, True)
            else:
                escribir_runs(parrafo, trozo[1:-1], negrita, True, codigo)
            continue
        run = parrafo.add_run(trozo)
        run.bold = negrita
        run.italic = cursiva
        if codigo:
            run.font.name = "Menlo"
            run.font.size = Pt(9.5)


def insertar_toc(doc):
    """Inserta un campo TOC nativo de Word. El texto no se calcula aquí:
    Word lo puebla al abrir el documento (o con F9). Es el mecanismo
    correcto — una lista escrita a mano se desincroniza al primer cambio."""
    parrafo = doc.add_paragraph()
    run = parrafo.add_run()

    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = r'TOC \o "1-3" \h \z \u'
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    marcador = OxmlElement("w:t")
    marcador.text = "Actualiza este campo (F9) para generar la tabla de contenido."
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")

    for elemento in (inicio, instruccion, separador, marcador, fin):
        run._r.append(elemento)


def es_separador(fila):
    """La fila |---|---| de Markdown no es datos: define la alineación."""
    return all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in fila if c.strip())


def celdas(linea):
    return [c.strip() for c in linea.strip().strip("|").split("|")]


MARGEN = Inches(0.9)
ANCHO_UTIL = Inches(8.5 - 2 * 0.9)  # carta menos márgenes

# Dentro de las tablas el código va un punto más pequeño: las rutas de
# campo del anexo ("estructura_grupos.filas_por_grupo.media") son tokens
# de 38 caracteres sin espacios, y a 9.5pt no caben en su columna, así
# que Word las parte por la mitad de la palabra.
PT_CODIGO_TABLA = Pt(8.5)


def anchos_por_contenido(encabezado, cuerpo, minimo=0.72):
    """Reparte el ancho disponible en proporción al texto de cada columna.

    Con columnas de ancho igual, una cabecera corta como "Marca" en una
    tabla de cuatro columnas se parte en dos líneas ("Mar/ca") mientras la
    columna de rutas largas queda apretada. El reparto proporcional evita
    las dos cosas."""
    n = len(encabezado)
    pesos = []
    for i in range(n):
        textos = [encabezado[i]] + [f[i] for f in cuerpo if i < len(f)]
        # El máximo, no la media ni un percentil: lo que se quiere evitar es
        # que una celda se parta por la mitad, y eso lo decide la celda más
        # larga, no la típica. En el anexo la columna "Valor" son 49 cifras
        # cortas y tres cadenas largas ("StratifiedGroupKFold"); promediar
        # deja esas tres rotas en cuatro líneas. El tope evita que una celda
        # excepcionalmente larga se lleve la tabla entera.
        mayor = max(len(t) for t in textos)
        pesos.append(max(min(mayor, 62), len(encabezado[i]) * 1.15, 1))

    # Las columnas que no llegan al mínimo se fijan ahí y se reparte el
    # resto entre las demás. Aplicar el mínimo y luego renormalizar sobre
    # el total lo deshace: era lo que dejaba "Marca" partido en "Marc/a".
    disponible = ANCHO_UTIL.inches
    fijas = {}
    libres = list(range(n))
    while True:
        resto = disponible - sum(fijas.values())
        suma_libres = sum(pesos[i] for i in libres)
        cortas = [i for i in libres if resto * pesos[i] / suma_libres < minimo]
        if not cortas:
            break
        for i in cortas:
            fijas[i] = minimo
            libres.remove(i)
        if not libres:
            break

    resto = disponible - sum(fijas.values())
    suma_libres = sum(pesos[i] for i in libres) or 1
    anchos = []
    for i in range(n):
        if i in fijas:
            anchos.append(Inches(fijas[i]))
        else:
            anchos.append(Inches(resto * pesos[i] / suma_libres))
    return anchos


def agregar_tabla(doc, filas):
    encabezado = filas[0]
    cuerpo = [f for f in filas[1:] if not es_separador(f)]
    tabla = doc.add_table(rows=1, cols=len(encabezado))
    tabla.style = "Table Grid"
    tabla.autofit = False

    anchos = anchos_por_contenido(encabezado, cuerpo)

    for i, texto in enumerate(encabezado):
        celda = tabla.rows[0].cells[i]
        celda.text = ""
        parrafo = celda.paragraphs[0]
        escribir_runs(parrafo, texto)
        for run in parrafo.runs:
            run.bold = True

    for fila in cuerpo:
        celdas_nuevas = tabla.add_row().cells
        for i, texto in enumerate(fila[: len(encabezado)]):
            celdas_nuevas[i].text = ""
            escribir_runs(celdas_nuevas[i].paragraphs[0], texto)

    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    if run.font.name == "Menlo":
                        run.font.size = PT_CODIGO_TABLA

    # Word ignora el ancho de columna si no se fija además en cada celda.
    for fila in tabla.rows:
        for i, celda in enumerate(fila.cells):
            if i < len(anchos):
                celda.width = anchos[i]

    doc.add_paragraph()
    return tabla


def bloques(lineas):
    """Agrupa las líneas en bloques lógicos. El Markdown viene con los
    párrafos partidos a 80 columnas, así que las líneas consecutivas de
    un mismo párrafo hay que volver a unirlas o el .docx sale con un
    salto de párrafo por cada línea del fuente."""
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        desnuda = linea.strip()

        if not desnuda:
            i += 1
            continue

        if re.fullmatch(r"-{3,}", desnuda):
            i += 1
            continue

        if desnuda.startswith("#"):
            nivel = len(desnuda) - len(desnuda.lstrip("#"))
            yield ("head", nivel, desnuda.lstrip("#").strip())
            i += 1
            continue

        if desnuda.startswith("|"):
            filas = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                filas.append(celdas(lineas[i]))
                i += 1
            yield ("table", 0, filas)
            continue

        if desnuda.startswith("> "):
            partes = []
            while i < len(lineas) and lineas[i].strip().startswith(">"):
                partes.append(lineas[i].strip().lstrip(">").strip())
                i += 1
            yield ("quote", 0, " ".join(partes).strip())
            continue

        marca_lista = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", linea)
        if marca_lista:
            sangria, marcador, resto = marca_lista.groups()
            partes = [resto]
            i += 1
            # las continuaciones de un ítem vienen indentadas y sin marca
            while i < len(lineas):
                siguiente = lineas[i]
                if not siguiente.strip():
                    break
                if re.match(r"^(\s*)([-*]|\d+\.)\s+", siguiente):
                    break
                if siguiente.startswith(("|", "#", ">")):
                    break
                partes.append(siguiente.strip())
                i += 1
            tipo = "ol" if marcador[0].isdigit() else "ul"
            yield (tipo, len(sangria), " ".join(partes))
            continue

        partes = []
        while i < len(lineas):
            siguiente = lineas[i]
            if not siguiente.strip():
                break
            if siguiente.strip().startswith(("|", "#", ">")):
                break
            if re.match(r"^(\s*)([-*]|\d+\.)\s+", siguiente):
                break
            if re.fullmatch(r"-{3,}", siguiente.strip()):
                break
            partes.append(siguiente.strip())
            i += 1
        yield ("p", 0, " ".join(partes))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--entrada", required=True)
    ap.add_argument("--salida", required=True)
    args = ap.parse_args()

    lineas = open(args.entrada, encoding="utf-8").read().split("\n")
    doc = Document()

    for seccion in doc.sections:
        seccion.left_margin = seccion.right_margin = MARGEN

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    titulo_puesto = False
    toc_puesto = False
    n_tablas = 0

    for tipo, nivel, contenido in bloques(lineas):
        if tipo == "head":
            grado = 0 if (nivel == 1 and not titulo_puesto) else min(nivel - 1, 4)
            titulo_puesto = titulo_puesto or grado == 0
            encabezado = doc.add_heading("", grado)
            escribir_runs(encabezado, contenido)
            continue

        if tipo == "quote":
            # El primer bloque de cita es la nota de lectura de la portada.
            # Después de ella va el índice, antes de la sección 1.
            parrafo = doc.add_paragraph(style="Intense Quote")
            escribir_runs(parrafo, contenido)
            if not toc_puesto:
                doc.add_page_break()
                # Rótulo con aspecto de encabezado pero sin serlo: si fuera
                # Heading 1, el propio índice se listaría como su primera
                # entrada.
                rotulo = doc.add_paragraph()
                run = rotulo.add_run("Contenido")
                run.bold = True
                run.font.size = Pt(18)
                insertar_toc(doc)
                doc.add_page_break()
                toc_puesto = True
            continue

        if tipo == "table":
            agregar_tabla(doc, contenido)
            n_tablas += 1
            continue

        if tipo in ("ul", "ol"):
            base = "List Bullet" if tipo == "ul" else "List Number"
            estilo_lista = base if nivel < 2 else f"{base} 2"
            parrafo = doc.add_paragraph(style=estilo_lista)
            escribir_runs(parrafo, contenido)
            continue

        parrafo = doc.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        escribir_runs(parrafo, contenido)

    doc.save(args.salida)
    print(f"Escrito: {args.salida} — {n_tablas} tablas")


if __name__ == "__main__":
    main()
